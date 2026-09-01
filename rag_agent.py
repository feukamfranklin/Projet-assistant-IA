"""Agent RAG LangChain (portage .py du notebook LangChain_RAG_Agent_v2.ipynb).

Ce module contient toute la logique de l'agent (vector store, chaine RAG, historique)
et est importe tel quel par API_REST.py : l'API ne fait qu'appeler ces fonctions.
"""

import base64
import json
import os
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from threading import Lock

import requests
from dotenv import load_dotenv, set_key
from docx import Document as WordDocument

from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openrouter import ChatOpenRouter
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
# ------------------------------------------------------------
# Configuration et initialisation du modele
# ------------------------------------------------------------
load_dotenv()
ENV_FILE_PATH = str(Path(__file__).resolve().parent / ".env")

POCKETBASE_URL = os.getenv('POCKETBASE_URL')
POCKETBASE_TOKEN = os.getenv("POCKETBASE_TOKEN")
# Compte utilise pour renouveler automatiquement le token (superuser ou utilisateur d'une collection auth).
POCKETBASE_AUTH_COLLECTION = os.getenv("POCKETBASE_AUTH_COLLECTION", "_superusers")
POCKETBASE_ADMIN_EMAIL = os.getenv("POCKETBASE_ADMIN_EMAIL")
POCKETBASE_ADMIN_PASSWORD = os.getenv("POCKETBASE_ADMIN_PASSWORD")
POCKETBASE_CLIENTS_COLLECTION = "clients"
POCKETBASE_COLLECTION = "Document_rag"
POCKETBASE_PROMPT = "system_prompt"
POCKETBASE_HISTORY_COLLECTION = "client_conversations"
VECTORSTORE_MANIFEST_RECORD_TITLE = "description_doc_sisma"
VECTORSTORE_MANIFEST_FILENAME = "doc_manifest.json"

_pb_token_lock = Lock()
_pb_token_exp = None  # epoch seconds d'expiration du token courant, None si inconnu


def _decode_jwt_exp(token):
    """Extrait le champ 'exp' d'un JWT PocketBase sans verifier la signature."""
    try:
        payload_b64 = token.split('.')[1]
        padding = '=' * (-len(payload_b64) % 4)
        payload = json.loads(base64.urlsafe_b64decode(payload_b64 + padding))
        return payload.get('exp')
    except Exception:
        return None


def _persist_pocketbase_token(token):
    """Ecrit le token renouvele dans .env pour qu'il survive a un redemarrage."""
    if not token:
        return
    try:
        set_key(ENV_FILE_PATH, "POCKETBASE_TOKEN", token)
        os.environ["POCKETBASE_TOKEN"] = token
    except Exception as exc:
        print(f"[pocketbase] echec ecriture du token dans .env: {exc}")


def _pocketbase_authenticate():
    """Recupere un nouveau token via identifiants admin/utilisateur."""
    global POCKETBASE_TOKEN, _pb_token_exp
    if not (POCKETBASE_URL and POCKETBASE_ADMIN_EMAIL and POCKETBASE_ADMIN_PASSWORD):
        return False
    try:
        response = requests.post(
            f"{POCKETBASE_URL}/api/collections/{POCKETBASE_AUTH_COLLECTION}/auth-with-password",
            json={"identity": POCKETBASE_ADMIN_EMAIL, "password": POCKETBASE_ADMIN_PASSWORD},
            timeout=10,
        )
        response.raise_for_status()
        data = response.json()
    except requests.RequestException as exc:
        print(f"[pocketbase] echec authentification ({POCKETBASE_AUTH_COLLECTION}): {exc}")
        return False
    POCKETBASE_TOKEN = data.get("token")
    _pb_token_exp = _decode_jwt_exp(POCKETBASE_TOKEN)
    _persist_pocketbase_token(POCKETBASE_TOKEN)
    return bool(POCKETBASE_TOKEN)


def _pocketbase_refresh_token():
    """Prolonge le token courant via auth-refresh, ou ré-authentifie en dernier recours."""
    global POCKETBASE_TOKEN, _pb_token_exp
    if POCKETBASE_URL and POCKETBASE_TOKEN:
        try:
            response = requests.post(
                f"{POCKETBASE_URL}/api/collections/{POCKETBASE_AUTH_COLLECTION}/auth-refresh",
                headers={'Authorization': f'Bearer {POCKETBASE_TOKEN}'},
                timeout=10,
            )
            response.raise_for_status()
            POCKETBASE_TOKEN = response.json().get("token")
            _pb_token_exp = _decode_jwt_exp(POCKETBASE_TOKEN)
            _persist_pocketbase_token(POCKETBASE_TOKEN)
            return True
        except requests.RequestException as exc:
            print(f"[pocketbase] echec refresh token: {exc}")
    return _pocketbase_authenticate()


def ensure_pocketbase_token():
    """Rafraichit POCKETBASE_TOKEN s'il est absent, expire ou proche de l'expiration."""
    global _pb_token_exp
    with _pb_token_lock:
        if _pb_token_exp is None and POCKETBASE_TOKEN:
            _pb_token_exp = _decode_jwt_exp(POCKETBASE_TOKEN)
        expires_soon = _pb_token_exp is not None and (_pb_token_exp - time.time()) < 300
        if not POCKETBASE_TOKEN or expires_soon:
            _pocketbase_refresh_token()


def pocketbase_headers():
    ensure_pocketbase_token()
    if not POCKETBASE_TOKEN:
        return {}
    return {'Authorization': f'Bearer {POCKETBASE_TOKEN}'}

if not os.getenv("OPENROUTER_API_KEY"):
    raise ValueError("Definissez OPENROUTER_API_KEY dans votre fichier .env avant de lancer l'agent.")

MODEL = "google/gemini-2.5-flash"
EMBEDDING_MODEL = "openai/text-embedding-3-small"

llm = ChatOpenRouter(
    model=MODEL,
    temperature=0.2,
)

embeddings = OpenAIEmbeddings(
    model=EMBEDDING_MODEL,
    openai_api_key=os.getenv("OPENROUTER_API_KEY"),
    openai_api_base="https://openrouter.ai/api/v1",
    headers={
        "Referer": "https://home.djamai.com",
        "X-OpenRouter-Title": "DjamAI Stage",
    },
)

ADDITIONAL_CONTEXT_PATH = "sys_Prompt.txt"  # repli local si la fiche PocketBase 'system_prompt' est indisponible
additional_context = ""

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150

# Historique de conversation stocke par session via PocketBase
MAX_HISTORY_TURNS = 4
MAX_OLDER_SUMMARY_CHARS = 800
MAX_STORED_TURNS = 30

# Verrou pour eviter les ecritures concurrentes sur l'historique
_history_lock = Lock()
_vectorstore_lock = Lock()
_vectorstore_cache = None
_vectorstore_signature = None


# ------------------------------------------------------------
# Construction / chargement du vector store FAISS
# ------------------------------------------------------------
def list_pocketbase_documents():
    """Retourne les enregistrements PocketBase contenant des fichiers."""
    if not POCKETBASE_URL:
        raise RuntimeError("POCKETBASE_URL est manquante.")

    response = requests.get(
        f"{POCKETBASE_URL}/api/collections/{POCKETBASE_COLLECTION}/records",
        headers=pocketbase_headers(),
        params={"perPage": 500},
        timeout=20,
    )
    response.raise_for_status()
    records = response.json().get("items", [])
    return [record for record in records if record.get("file")]


def build_pocketbase_manifest(records):
    """Construit un manifest stable pour détecter les changements PocketBase."""
    return sorted(
        [
            {
                "id": record.get("id"),
                "file": record.get("file"),
                "updated": record.get("updated"),
            }
            for record in records
        ],
        key=lambda item: item["id"] or "",
    )


def download_pocketbase_documents(records):
    """Télécharge temporairement les fichiers PocketBase pour les loaders LangChain."""
    temp_dir = tempfile.TemporaryDirectory(prefix="rag_pocketbase_")
    downloaded_paths = []

    for record in records:
        filenames = record["file"] if isinstance(record["file"], list) else [record["file"]]
        for filename in filenames:
            response = requests.get(
                f"{POCKETBASE_URL}/api/files/{POCKETBASE_COLLECTION}/{record['id']}/{filename}",
                headers=pocketbase_headers(),
                timeout=30,
            )
            response.raise_for_status()
            destination = Path(temp_dir.name) / f"{record['id']}_{Path(filename).name}"
            destination.write_bytes(response.content)
            downloaded_paths.append(str(destination))

    return temp_dir, downloaded_paths


def load_documents_from_pocketbase(records):
    """Charge les fichiers PocketBase avec les loaders documentaires existants."""
    temp_dir, downloaded_paths = download_pocketbase_documents(records)
    try:
        documents = []
        for file_path in downloaded_paths:
            ext = Path(file_path).suffix.lower()
            if ext in {".txt", ".md"}:
                documents.extend(TextLoader(file_path, encoding="utf-8").load())
            elif ext == ".pdf":
                documents.extend(PyPDFLoader(file_path).load())
            elif ext == ".docx":
                documents.extend(load_docx(file_path))
        if not documents:
            raise FileNotFoundError(
                f"Aucun fichier exploitable dans la collection '{POCKETBASE_COLLECTION}'."
            )
        return documents
    finally:
        temp_dir.cleanup()


def load_docx(path: str):
    """Charge un fichier DOCX et retourne le texte en document LangChain."""
    word_doc = WordDocument(path)
    text = "\n".join(p.text for p in word_doc.paragraphs if p.text and p.text.strip())
    if not text.strip():
        return []
    return [Document(page_content=text, metadata={"source": path})]


def _document_records_url() -> str:
    if not POCKETBASE_URL:
        raise RuntimeError("POCKETBASE_URL est manquante.")
    return f"{POCKETBASE_URL}/api/collections/{POCKETBASE_COLLECTION}/records"


def _find_document_record_by_title(title: str):
    response = requests.get(
        _document_records_url(),
        headers=pocketbase_headers(),
        params={"perPage": 500},
        timeout=20,
    )
    response.raise_for_status()

    for record in response.json().get("items", []):
        if record.get("titre") == title:
            return record
    return None


def _download_manifest_from_document_record(record):
    files = record.get("file")
    filenames = files if isinstance(files, list) else ([files] if files else [])
    json_name = next((name for name in filenames if str(name).lower().endswith(".json")), None)
    if not json_name:
        return None

    response = requests.get(
        f"{POCKETBASE_URL}/api/files/{POCKETBASE_COLLECTION}/{record['id']}/{json_name}",
        headers=pocketbase_headers(),
        timeout=20,
    )
    response.raise_for_status()
    try:
        return json.loads(response.content.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError):
        return None


def _download_text_file_from_document_record(record):
    files = record.get("file")
    filenames = files if isinstance(files, list) else ([files] if files else [])
    txt_name = next((name for name in filenames if str(name).lower().endswith((".txt", ".md"))), None)
    if not txt_name:
        return None

    response = requests.get(
        f"{POCKETBASE_URL}/api/files/{POCKETBASE_COLLECTION}/{record['id']}/{txt_name}",
        headers=pocketbase_headers(),
        timeout=20,
    )
    response.raise_for_status()
    try:
        return response.content.decode("utf-8")
    except UnicodeDecodeError:
        return None


def load_system_prompt():
    """Charge le system prompt depuis la fiche 'system_prompt' de Document_rag, avec repli sur le fichier local."""
    global additional_context
    try:
        record = _find_document_record_by_title(POCKETBASE_PROMPT)
        if record:
            text = _download_text_file_from_document_record(record)
            if text:
                additional_context = text
                return additional_context
    except requests.RequestException as exc:
        print(f"[pocketbase] echec chargement du system_prompt: {exc}")
    return additional_context


def _load_persisted_vectorstore_signature(record=None):
    """Charge la signature vectorstore depuis le fichier manifest de Document_rag."""
    if record is None:
        record = _find_document_record_by_title(VECTORSTORE_MANIFEST_RECORD_TITLE)
    if not record:
        return None

    manifest = _download_manifest_from_document_record(record)
    if not isinstance(manifest, list):
        return None
    return json.dumps(manifest, sort_keys=True, ensure_ascii=True)


def _save_vectorstore_manifest(current_manifest, current_signature, record=None):
    """Sauvegarde/replace le fichier manifest JSON dans Document_rag."""
    manifest_bytes = json.dumps(current_manifest, ensure_ascii=False, indent=2).encode("utf-8")
    data = {"titre": VECTORSTORE_MANIFEST_RECORD_TITLE}
    files = {"file": (VECTORSTORE_MANIFEST_FILENAME, manifest_bytes, "application/json")}

    if record is None:
        record = _find_document_record_by_title(VECTORSTORE_MANIFEST_RECORD_TITLE)
    if record:
        response = requests.patch(
            f"{_document_records_url()}/{record['id']}",
            headers=pocketbase_headers(),
            data=data,
            files=files,
            timeout=20,
        )
    else:
        response = requests.post(
            _document_records_url(),
            headers=pocketbase_headers(),
            data=data,
            files=files,
            timeout=20,
        )
    response.raise_for_status()


def load_or_build_vectorstore():
    """Construit un vector store en memoire depuis PocketBase et le met en cache."""
    global _vectorstore_cache, _vectorstore_signature

    records = list_pocketbase_documents()
    if not records:
        raise FileNotFoundError(
            f"Aucun fichier trouve dans la collection '{POCKETBASE_COLLECTION}'."
        )
    current_manifest = build_pocketbase_manifest(records)
    current_signature = json.dumps(
        current_manifest,
        sort_keys=True,
        ensure_ascii=True,
    )

    with _vectorstore_lock:
        manifest_record = None
        if _vectorstore_signature is None:
            manifest_record = _find_document_record_by_title(VECTORSTORE_MANIFEST_RECORD_TITLE)
            _vectorstore_signature = _load_persisted_vectorstore_signature(manifest_record)

        if _vectorstore_cache is not None and current_signature == _vectorstore_signature:
            return _vectorstore_cache

        documents = load_documents_from_pocketbase(records)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
        )
        chunks = text_splitter.split_documents(documents)

        _vectorstore_cache = FAISS.from_documents(chunks, embeddings)
        _save_vectorstore_manifest(current_manifest, current_signature, manifest_record)
        _vectorstore_signature = current_signature
        return _vectorstore_cache


def get_retriever():
    """Cree un retriever qui extrait les meilleurs passages du vector store."""
    vectorstore = load_or_build_vectorstore()
    return vectorstore.as_retriever(search_kwargs={"k": 30})


# ------------------------------------------------------------
# Chaine RAG
# ------------------------------------------------------------
rag_prompt_template = """Vous etes un assistant de sisma fournissant des informations sur les services de l'etablissement.
Utilisez les informations extraites de la base de donnees comme source prioritaire pour repondre et utilise aussi {additional_context} pour mieux connaitre ce que tu dois faire.
Quand la question demande une liste (ex: filieres), retournez la liste de maniere structuree.

Historique recent de conversation :
{conversation_history}

Contexte :
{context}

Question : {question}
"""

prompt = ChatPromptTemplate.from_template(rag_prompt_template)
rag_chain = prompt | llm | StrOutputParser()

retriever = None  # initialise via init_retriever(), appele au demarrage de l'API


def init_retriever():
    """Initialise (ou reinitialise) le retriever et le system prompt. A appeler au demarrage de l'API."""
    global retriever
    retriever = get_retriever()
    load_system_prompt()
    return retriever


def format_docs(docs):
    """Joindre les morceaux de document recuperes dans une seule chaine de contexte."""
    return "\n\n".join(doc.page_content for doc in docs)


# ------------------------------------------------------------
# Historique de conversation (PocketBase)
# ------------------------------------------------------------
def _history_records_url() -> str:
    if not POCKETBASE_URL:
        raise RuntimeError("POCKETBASE_URL est manquante.")
    return f"{POCKETBASE_URL}/api/collections/{POCKETBASE_HISTORY_COLLECTION}/records"


def _clients_records_url() -> str:
    if not POCKETBASE_URL:
        raise RuntimeError("POCKETBASE_URL est manquante.")
    return f"{POCKETBASE_URL}/api/collections/{POCKETBASE_CLIENTS_COLLECTION}/records"


def _pocketbase_datetime_now() -> str:
    # Format date attendu par PocketBase: UTC avec precision millisecondes.
    return datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S.%f")[:-3] + "Z"


def _find_client_by_phone(num_tel: str):
    response = requests.get(
        _clients_records_url(),
        headers=pocketbase_headers(),
        params={"filter": f'num_tel = "{num_tel}"', "perPage": 1},
        timeout=20,
    )
    response.raise_for_status()
    return (response.json().get("items") or [None])[0]


def _get_or_create_client(num_tel: str):
    client = _find_client_by_phone(num_tel)
    if client:
        return client

    response = requests.post(
        _clients_records_url(),
        headers=pocketbase_headers(),
        json={"num_tel": num_tel},
        timeout=20,
    )
    response.raise_for_status()
    return response.json()


def _list_history_records():
    records = []
    page = 1

    while True:
        params = {"page": page, "perPage": 200, "sort": "-updated"}
        response = requests.get(
            _history_records_url(),
            headers=pocketbase_headers(),
            params=params,
            timeout=20,
        )
        if response.status_code == 400:
            # Certaines versions/configurations PocketBase refusent ce tri sur cette collection.
            params.pop("sort", None)
            response = requests.get(
                _history_records_url(),
                headers=pocketbase_headers(),
                params=params,
                timeout=20,
            )
        response.raise_for_status()
        payload = response.json()
        items = payload.get("items", [])
        records.extend(items)

        total_pages = payload.get("totalPages", 1)
        if page >= total_pages:
            break
        page += 1

    return records


def _extract_history_payload(record):
    data = record.get("json_file")
    return data if isinstance(data, list) else []


def _find_history_record_by_client_id(client_id: str):
    for record in _list_history_records():
        if record.get("client_id") == client_id:
            return record, _extract_history_payload(record)
    return None, []


def load_conversation_history(num_tel: str):
    """Charge l'historique d'un client identifie par son numero de telephone."""
    try:
        client = _find_client_by_phone(num_tel)
        if not client:
            return []
        _, history = _find_history_record_by_client_id(client["id"])
        return history if isinstance(history, list) else []
    except requests.RequestException:
        return []


def save_conversation_history(num_tel: str, history):
    """Sauvegarde l'historique du client dans PocketBase en limitant sa taille."""
    client = _get_or_create_client(num_tel)
    trimmed = history[-MAX_STORED_TURNS:] if len(history) > MAX_STORED_TURNS else history
    payload = {
        "date_time": _pocketbase_datetime_now(),
        "client_id": client["id"],
        "json_file": trimmed,
    }

    record, _ = _find_history_record_by_client_id(client["id"])
    if record:
        response = requests.patch(
            f"{_history_records_url()}/{record['id']}",
            headers=pocketbase_headers(),
            json=payload,
            timeout=20,
        )
    else:
        response = requests.post(
            _history_records_url(),
            headers=pocketbase_headers(),
            json=payload,
            timeout=20,
        )
    response.raise_for_status()


def clear_conversation_history(num_tel: str):
    """Supprime l'historique d'un client depuis PocketBase."""
    client = _find_client_by_phone(num_tel)
    if not client:
        return
    record, _ = _find_history_record_by_client_id(client["id"])
    if not record:
        return

    response = requests.delete(
        f"{_history_records_url()}/{record['id']}",
        headers=pocketbase_headers(),
        timeout=20,
    )
    response.raise_for_status()


def format_conversation_history(history, max_turns: int = MAX_HISTORY_TURNS, max_summary_chars: int = MAX_OLDER_SUMMARY_CHARS):
    """Construit un historique compact pour le prompt du LLM."""
    if not history:
        return "Aucun historique disponible."

    if len(history) <= max_turns:
        lines = []
        for turn in history:
            user_text = (turn.get("user") or "").strip()
            assistant_text = (turn.get("assistant") or "").strip()
            if user_text:
                lines.append(f"Utilisateur: {user_text}")
            if assistant_text:
                lines.append(f"Assistant: {assistant_text}")
        return "\n".join(lines)

    recent = history[-max_turns:]
    older = history[:-max_turns]

    older_parts = []
    for turn in older:
        user_text = (turn.get("user") or "").strip()
        assistant_text = (turn.get("assistant") or "").strip()
        if user_text or assistant_text:
            older_parts.append(f"Q: {user_text} | R: {assistant_text}")

    older_summary = " ".join(older_parts)[:max_summary_chars]

    recent_lines = []
    for turn in recent:
        user_text = (turn.get("user") or "").strip()
        assistant_text = (turn.get("assistant") or "").strip()
        if user_text:
            recent_lines.append(f"Utilisateur: {user_text}")
        if assistant_text:
            recent_lines.append(f"Assistant: {assistant_text}")

    recent_text = "\n".join(recent_lines)
    return f"Resume des echanges precedents : {older_summary}\n\nHistorique recent :\n{recent_text}"


def ask_rag(question: str, num_tel: str) -> tuple[str, str]:
    """Pose une question a l'agent RAG pour un numero donne."""
    if retriever is None:
        init_retriever()

    with _history_lock:
        history = load_conversation_history(num_tel)
        history_text = format_conversation_history(history)

        retrieved = retriever.invoke(question)
        context = format_docs(retrieved)

        answer = rag_chain.invoke({
            "context": context,
            "question": question,
            "additional_context": additional_context,
            "conversation_history": history_text,
        })

        history.append({"user": question.strip(), "assistant": answer.strip()})
        save_conversation_history(num_tel, history)

    return answer, context
